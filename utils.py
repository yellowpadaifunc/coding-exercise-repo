"""
Utility functions for the project.
This module contains helper functions that can be used across different parts of the application.
"""

from copy import deepcopy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def insert_before_with_numbering(source_para, text):
    # 1. Insert the new paragraph just before the source paragraph
    new_para = source_para.insert_paragraph_before(text)
    
    # 2. Copy over the same style
    new_para.style = source_para.style

    # 3. Copy over the numbering properties (numId & ilvl)
    src_p = source_para._p
    tgt_p = new_para._p

    # Ensure tgt_p has a <w:pPr>
    pPr = tgt_p.get_or_add_pPr()

    # Look for the <w:numPr> element in the source's pPr
    numPr = None
    if src_p.pPr is not None:
        numPr = src_p.pPr.find(qn("w:numPr"))

    # If found, deep‑copy it and append to the new paragraph
    if numPr is not None:
        pPr.append(deepcopy(numPr))

    return new_para

def get_numbering_props(paragraph):
    """
    Returns (numId, ilvl) for a numbered/list paragraph,
    or (None, None) if it's not part of any list.
    """
    pPr = paragraph._p.pPr
    if pPr is None:
        return None, None
    
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None, None
    
    numId = numPr.find(qn("w:numId")).get(qn("w:val"))
    ilvl = numPr.find(qn("w:ilvl")).get(qn("w:val"))
    return numId, ilvl

def has_quotes(text):
    text = text.strip()
    # Only check if the first word/phrase is quoted (straight or curly)
    return (
        (text.startswith('"') and '"' in text[1:]) or
        (text.startswith('“') and '”' in text[1:])
    )

def is_main_heading(para):
    text = para.text.strip()
    # Heuristic: main headings end with a period, are short, and not quoted
    return (
        text.endswith('.') and
        len(text.split()) < 10 and
        not (text.startswith('"') or text.startswith('"'))
    )

def find_next_main_heading_index(document, start_index):
    for i in range(start_index + 1, len(document.paragraphs)):
        para = document.paragraphs[i]
        if is_main_heading(para):
            return i
    return None

def detect_clause_heading_style(document, section_heading_index=None):
    """
    Analyze the first few paragraphs (or those in a section) to infer clause heading style.
    Returns a dict: {"bold": bool, "underline": bool, "quotes": bool}
    If section_heading_index is provided, only scan that section.
    """
    if section_heading_index is not None:
        next_section_index = find_next_main_heading_index(document, section_heading_index)
        if next_section_index is None:
            end_index = len(document.paragraphs)
        else:
            end_index = next_section_index
        para_range = range(section_heading_index + 1, end_index)
    else:
        para_range = range(len(document.paragraphs))

    for i in para_range:
        para = document.paragraphs[i]
        text = para.text.strip()

        if not text:
            continue
        # Heuristic: clause headings are often short and styled
        if para.runs and len(text) < 60:
            bold = any(run.bold for run in para.runs)
            underline = any(run.underline for run in para.runs)
            quotes = has_quotes(text)
            # Only consider if at least bold or underline or quotes
            if bold or underline or quotes:
                return {"bold": bold, "underline": underline, "quotes": quotes}
    # Default fallback
    return {"bold": True, "underline": False, "quotes": False}

def format_clause_heading(heading_text, style_info, document=None):
    """
    Returns a formatted docx paragraph for the heading, or a dict for insertion.
    If document is provided, adds the paragraph to the document and returns it.
    """
    if style_info.get("quotes"):
        heading_text = f'"{heading_text}"'
    if document is not None:
        para = document.add_paragraph()
        run = para.add_run(heading_text)
        run.bold = style_info.get("bold", False)
        run.underline = style_info.get("underline", False)
        return para
    else:
        return {
            "text": heading_text,
            "bold": style_info.get("bold", False),
            "underline": style_info.get("underline", False)
        }

def parse_clauses(document):
    """
    Returns a list of (heading_paragraph, heading_text, index) for all detected clause headings.
    """
    style_info = detect_clause_heading_style(document)

    clauses = []
    for i, para in enumerate(document.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        # Heuristic: matches style
        is_heading = False
        if style_info.get("quotes") and text.startswith('"') and text.endswith('"'):
            is_heading = True
        elif style_info.get("bold") and any(run.bold for run in para.runs):
            if style_info.get("underline"):
                if any(run.underline for run in para.runs):
                    is_heading = True
            else:
                is_heading = True
        if is_heading:
            clauses.append((para, text, i))
    return clauses

def insert_clause(document, clause_heading, clause_body, position=None, style_info=None):
    """
    Inserts a new clause (heading + body) at the specified position (index in paragraphs).
    If position is None, appends at the end. style_info can be provided or auto-detected.
    Returns the new heading and body paragraphs.
    """
    if style_info is None:
        style_info = detect_clause_heading_style(document)
    if position is None or position >= len(document.paragraphs):
        heading_para = format_clause_heading(clause_heading, style_info, document)
        body_para = document.add_paragraph(clause_body)
        return heading_para, body_para
    else:
        # Insert at position: need to use _element API
        paras = document.paragraphs
        ref_para = paras[position]
        # Insert heading before ref_para
        heading_para = ref_para.insert_paragraph_before("")
        run = heading_para.add_run(clause_heading)
        run.bold = style_info.get("bold", False)
        run.underline = style_info.get("underline", False)
        if style_info.get("quotes"):
            run.text = f'"{clause_heading}"'
        # Insert body after heading
        body_para = heading_para.insert_paragraph_after(clause_body)
        return heading_para, body_para

def find_clause_index_by_heading(document, heading, case_insensitive=True, ignore_quotes=True):
    """
    Returns the index of the clause heading that matches the given heading string.
    If not found, returns None.
    - case_insensitive: if True, ignores case when matching
    - ignore_quotes: if True, strips quotes (curly or straight) from both document and input heading before matching
    - Normalizes curly and straight quotes for robust matching
    """
    def normalize(text):
        t = text.strip()
        # Remove both straight and curly quotes from start and end
        QUOTES = ['"', '\u201c', '\u201d', '\u201e', '\u201f', '\u2033', '\u2036']
        while t and (t[0] in QUOTES):
            t = t[1:]
        while t and (t[-1] in QUOTES):
            t = t[:-1]
        return t.lower() if case_insensitive else t

    target = normalize(heading)
    for i, para in enumerate(document.paragraphs):
        para_text = normalize(para.text)
        if para_text == target:
            return i
    return None

def insert_paragraph_after(paragraph, text):
    """
    Inserts a new paragraph after the given paragraph and returns the new paragraph.
    """
    p = paragraph._p
    new_p = OxmlElement('w:p')
    p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para

def find_next_heading_index(document, start_index, style_info=None):
    """
    Returns the index of the next heading-style paragraph after start_index.
    """
    if style_info is None:
        style_info = detect_clause_heading_style(document)
    for i in range(start_index + 1, len(document.paragraphs)):
        para = document.paragraphs[i]
        text = para.text.strip()
        if not text:
            continue
        is_heading = False
        if style_info.get("quotes") and text.startswith('"') and text.endswith('"'):
            is_heading = True
        elif style_info.get("bold") and any(run.bold for run in para.runs):
            if style_info.get("underline"):
                if any(run.underline for run in para.runs):
                    is_heading = True
            else:
                is_heading = True
        if is_heading:
            return i
    return None

def apply_numbering_from_template(new_para, template_para):
    """
    Copies numbering properties from template_para to new_para, if present.
    """
    numId, ilvl = get_numbering_props(template_para)
    if numId and ilvl:
        pPr = new_para._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        numIdEl = OxmlElement('w:numId')
        numIdEl.set(qn('w:val'), numId)
        ilvlEl = OxmlElement('w:ilvl')
        ilvlEl.set(qn('w:val'), ilvl)
        numPr.append(ilvlEl)
        numPr.append(numIdEl)
        pPr.append(numPr)

def detect_quote_style(text, default_to_curly=True):
    """
    Detects the quote style used in the given text.
    Returns a tuple (open_quote, close_quote) or (None, None) if no quotes.
    If default_to_curly is True and no quotes are detected, returns curly quotes.
    """
    text = text.strip()
    if text.startswith('"') and text.endswith('"'):
        return '"', '"'

    if default_to_curly:
        return '“', '”'
    return None, None

def find_first_heading_with_quotes(document, start_index):
    """
    Finds the first heading-style paragraph after start_index that uses quotes.
    Returns (open_quote, close_quote) or (None, None) if not found.
    """
    for i in range(start_index + 1, len(document.paragraphs)):
        para = document.paragraphs[i]
        text = para.text.strip()
        open_quote, close_quote = detect_quote_style(text)
        if open_quote and close_quote:
            return open_quote, close_quote
    return None, None

def section_clause_headings_use_quotes(document, section_heading_index, style_info=None):
    """
    Returns True if the majority of clause headings under the section use quotes, False otherwise.
    Only scans until the next main heading (same detection as find_next_heading_index).
    """
    if style_info is None:
        style_info = detect_clause_heading_style(document, section_heading_index)
    next_section_index = find_next_main_heading_index(document, section_heading_index)
    if next_section_index is None:
        end_index = len(document.paragraphs)
    else:
        end_index = next_section_index
                
    quote_count = 0
    total = 0
    for i in range(section_heading_index + 1, end_index):
        para = document.paragraphs[i]
        text = para.text.strip()
        if not text:
            continue
        is_heading = False
        if style_info.get("quotes") and (text.startswith('"') and text.endswith('"') or text.startswith('"') and text.endswith('"')):
            is_heading = True
        elif style_info.get("bold") and any(run.bold for run in para.runs):
            if style_info.get("underline"):
                if any(run.underline for run in para.runs):
                    is_heading = True
            else:
                is_heading = True
        if is_heading:
            total += 1
            # Check for quotes (straight or curly)
            if (text.startswith('"') and text.endswith('"')) or (text.startswith('"') and text.endswith('"')):
                quote_count += 1
    if total == 0:
        return style_info.get("quotes", False)
    return quote_count > (total // 2)

def insert_styled_numbered_clause_after_heading(document, after_heading, heading_text, definition_text, after_clause=None):
    """
    Inserts a new clause (heading + definition) after the specified heading.
    If after_clause is provided, inserts after that clause heading within the section; otherwise, inserts as the first clause under the heading.
    The new clause will match the style and numbering of the next clause after the insertion point.
    Normalizes curly and straight quotes for robust matching.
    Matches the quote style (curly or straight) of the section's clause headings.
    """
    
    index = find_clause_index_by_heading(document, after_heading)
    if index is None:
        raise ValueError(f"Heading '{after_heading}' not found in document.")
    style_info = detect_clause_heading_style(document, index)

    # Determine insertion point
    insert_after_index = index
    if after_clause is not None:
        # Search for after_clause after after_heading
        style_info = detect_clause_heading_style(document, index)
        def normalize(text):
            t = text.strip()
            # Remove both straight and curly quotes from start and end
            QUOTES = ['"', '\u201c', '\u201d', '\u201e', '\u201f', '\u2033', '\u2036']
            while t and (t[0] in QUOTES):
                t = t[1:]
            while t and (t[-1] in QUOTES):
                t = t[:-1]
            return t.lower()
        norm_target = normalize(after_clause)
        for i in range(index + 1, len(document.paragraphs)):
            para = document.paragraphs[i]
            text = para.text.strip()
            is_heading = False
            if style_info.get("quotes") and (text.startswith('"') and text.endswith('"') or text.startswith('"') and text.endswith('"')):
                is_heading = True
            elif style_info.get("bold") and any(run.bold for run in para.runs):
                if style_info.get("underline"):
                    if any(run.underline for run in para.runs):
                        is_heading = True
                else:
                    is_heading = True
            if is_heading:
                norm_para = normalize(text)
                # Match if after_clause matches the start of the paragraph (after normalization)
                if norm_para.startswith(norm_target):
                    insert_after_index = i
                    break
        else:
            raise ValueError(f"after_clause '{after_clause}' not found under heading '{after_heading}'.")
    # If after_clause is not provided, insert immediately after after_heading
    else:
        insert_after_index = index

    # Dynamically determine style for this section
    section_style = section_clause_headings_style(document, index, style_info)

    use_quotes = section_style["use_quotes"]
    use_bold = section_style["use_bold"]
    use_italic = section_style["use_italic"]
    use_underline = section_style["use_underline"]

    # Find the next heading after the insertion point for style/numbering
    next_heading_index = find_next_heading_index(document, insert_after_index, style_info)
    if next_heading_index is not None:
        template_para = document.paragraphs[next_heading_index]
    else:
        template_para = document.paragraphs[-1]
    # Detect quote style from template_para
    template_text = template_para.text.strip()
    open_quote, close_quote = detect_quote_style(template_text)
    # If template has no quotes, look for first heading with quotes in section
    if not (open_quote and close_quote):
        open_quote, close_quote = find_first_heading_with_quotes(document, insert_after_index)
    # If still not found, default to curly quotes
    if not (open_quote and close_quote):
        open_quote, close_quote = detect_quote_style('', default_to_curly=True)
    # Build the heading with quotes as separate runs (quotes not bolded, heading styled)
    new_para = insert_before_with_numbering(template_para, "")
    if use_quotes and open_quote and close_quote:
        quote_run = new_para.add_run(open_quote)
        heading_run = new_para.add_run(heading_text)
        heading_run.bold = use_bold
        heading_run.italic = use_italic
        heading_run.underline = use_underline
        quote_run2 = new_para.add_run(close_quote)
    else:
        heading_run = new_para.add_run(heading_text)
        heading_run.bold = use_bold
        heading_run.italic = use_italic
        heading_run.underline = use_underline
    new_para.add_run(" " + definition_text)
    # Copy numbering from template
    apply_numbering_from_template(new_para, template_para)
    return new_para

def is_clause_heading(para, style_info):
    text = para.text.strip()
    if not text:
        return False
    if style_info.get("quotes") and has_quotes(text):
        return True
    if style_info.get("bold") and any(run.bold for run in para.runs):
        if style_info.get("underline"):
            if any(run.underline for run in para.runs):
                return True
        else:
            return True
    return False

def section_clause_headings_style(document, section_heading_index, style_info=None):
    """
    Returns a dict with use_quotes, use_bold, use_italic, use_underline for clause headings in the section.
    Each is True if the majority of clause headings in the section use that style.
    Only considers paragraphs that are actually clause headings.
    """
    if style_info is None:
        style_info = detect_clause_heading_style(document, section_heading_index)
    next_section_index = find_next_main_heading_index(document, section_heading_index)
    if next_section_index is None:
        end_index = len(document.paragraphs)
    else:
        end_index = next_section_index

    quote_count = 0
    bold_count = 0
    italic_count = 0
    underline_count = 0
    total = 0
    for i in range(section_heading_index + 1, end_index):
        para = document.paragraphs[i]
        if not is_clause_heading(para, style_info):
            continue
        text = para.text.strip()
        # Quotes: use has_quotes helper
        has_quotes_val = has_quotes(text)
        # Bold, Italic, Underline: any run in para
        has_bold = any(run.bold for run in para.runs)
        has_italic = any(
            (run.italic is True) or
            (getattr(getattr(run, 'style', None), 'font', None) and getattr(run.style.font, 'italic', None) is True)
            for run in para.runs
        )
        if not has_italic:
            # Check paragraph style and its base styles
            style = para.style
            while style is not None:
                if getattr(getattr(style, 'font', None), 'italic', None) is True:
                    has_italic = True
                    break
                style = getattr(style, 'base_style', None)
        # Fallback: check for <w:i/> in the XML
        if not has_italic and para._p.xpath('.//w:i'):
            has_italic = True
        has_underline = any(run.underline for run in para.runs)
        total += 1
        if has_quotes_val:
            quote_count += 1
        if has_bold:
            bold_count += 1
        if has_italic:
            italic_count += 1
        if has_underline:
            underline_count += 1

    if total == 0:
        return {
            "use_quotes": style_info.get("quotes", False),
            "use_bold": style_info.get("bold", False),
            "use_italic": False,
            "use_underline": style_info.get("underline", False)
        }
    return {
        "use_quotes": quote_count > (total // 2),
        "use_bold": bold_count > (total // 2),
        "use_italic": italic_count > (total // 2),
        "use_underline": underline_count > (total // 2)
    }

def renumber_clauses_preserve_formatting(document):
    """
    Renumber all top-level numbered clauses in the document sequentially, updating only the leading number and period in each paragraph, preserving all formatting.
    Assumes clause headings start with a number or 'XXX' and a period (e.g., '1. Heading', 'XXX. Heading').
    """
    clause_number = 1
    clause_heading_pattern = re.compile(r'^(\d+|XXX)\.\s+')
    for para in document.paragraphs:
        match = clause_heading_pattern.match(para.text.strip())
        if match:
            # Concatenate all run texts to find the number/period span
            total = ''.join(run.text for run in para.runs)
            m = clause_heading_pattern.match(total)
            if m:
                start, end = m.span()
                chars_seen = 0
                for run in para.runs:
                    run_len = len(run.text)
                    if chars_seen < end and chars_seen + run_len > 0:
                        rel_start = max(0, start - chars_seen)
                        rel_end = min(run_len, end - chars_seen)
                        run.text = (run.text[:rel_start] + f"{clause_number}. " + run.text[rel_end:])
                        break
                    chars_seen += run_len
            clause_number += 1

def add_heading_runs(paragraph, number_part, heading_part, period_part, src_run, style_info):
    """
    Adds three runs to the paragraph: number, heading, period.
    - number_part: e.g., '11. '
    - heading_part: e.g., 'Residuals'
    - period_part: e.g., '. '
    The number is bold (not underlined), heading is bold/italic/underlined, period is normal.
    Font name/size/color are copied from src_run if present.
    """
    def copy_font(dst_run, src_run, bold=None, italic=None, underline=None):
        if src_run:
            if src_run.font.name:
                dst_run.font.name = src_run.font.name
            if src_run.font.size:
                dst_run.font.size = src_run.font.size
            if src_run.font.color and src_run.font.color.rgb:
                dst_run.font.color.rgb = src_run.font.color.rgb
        if bold is not None:
            dst_run.bold = bold
        if italic is not None:
            dst_run.italic = italic
        if underline is not None:
            dst_run.underline = underline
    # Number run: bold, not underlined
    number_run = paragraph.add_run(number_part)
    copy_font(number_run, src_run, bold=True, italic=False, underline=False)
    # Heading run: bold, italic, underlined as in src_run/style_info
    heading_run = paragraph.add_run(heading_part)
    copy_font(
        heading_run,
        src_run,
        bold=src_run.bold if src_run and src_run.bold is not None else style_info.get("bold", False),
        italic=src_run.italic if src_run and src_run.italic is not None else style_info.get("italic", False),
        underline=src_run.underline if src_run and src_run.underline is not None else style_info.get("underline", False),
    )
    # Period run: normal
    period_run = paragraph.add_run(period_part)
    copy_font(period_run, src_run, bold=False, italic=False, underline=False)
    return number_run, heading_run, period_run

def insert_styled_clause_after_clause(document, after_clause_heading, new_clause_heading, new_clause_body):
    """
    Inserts a new clause (as a single paragraph: number, heading, and body together) after the specified clause heading (by text, not by section),
    and renumbers all top-level numbered clauses in the document to maintain sequential numbering.
    - after_clause_heading: the text of the clause heading after which to insert (e.g., '10. Prohibition on Use of Open AI Systems.' or 'Prohibition on Use of Open AI Systems.')
    - new_clause_heading: the heading text for the new clause (without number)
    - new_clause_body: the body text for the new clause
    Returns the new clause paragraph.
    """
    def normalize_heading(text):
        # Remove leading number and period, then strip and lowercase
        return re.sub(r'^\d+\.\s*', '', text.strip()).lower()

    def extract_clause_heading(text):
        """
        Extracts the clause heading (number + title + period) from a paragraph.
        E.g., '10. Prohibition on Use of Open AI Systems. Notwithstanding ...' -> '10. Prohibition on Use of Open AI Systems.'
        """
        match = re.match(r'^(\d+\.\s*[^.]+\.)', text.strip())
        if match:
            return match.group(1)
        return text.strip()

    norm_target = normalize_heading(after_clause_heading)
    after_index = None
    for i, para in enumerate(document.paragraphs):
        para_heading = extract_clause_heading(para.text)
        if normalize_heading(para_heading) == norm_target:
            after_index = i
            break
    if after_index is None:
        raise ValueError(f"Clause heading '{after_clause_heading}' not found in document.")

    # Insert a blank paragraph (line break) before the new clause
    after_para = document.paragraphs[after_index]
    blank_para = insert_paragraph_after(after_para, "")
    new_para = insert_paragraph_after(blank_para, "")

    # Robustly split number, heading, and period
    full_heading = f"XXX. {new_clause_heading}."
    first_dot_space = full_heading.find('. ')
    last_dot = full_heading.rfind('.')
    if first_dot_space != -1 and last_dot != -1 and last_dot > first_dot_space:
        number_part = full_heading[:first_dot_space+2]  # include '. '
        heading_part = full_heading[first_dot_space+2:last_dot]
        period_part = full_heading[last_dot] + " "
    else:
        number_part = ""
        heading_part = full_heading.strip()
        period_part = " "

    style_info = detect_clause_heading_style(document)
    src_run = after_para.runs[0] if after_para.runs else None
    # Use the helper to add heading runs
    add_heading_runs(new_para, number_part, heading_part, period_part, src_run, style_info)
    # Add the body (not bold/underlined)
    body_run = new_para.add_run(new_clause_body)
    if src_run:
        if src_run.font.name:
            body_run.font.name = src_run.font.name
        if src_run.font.size:
            body_run.font.size = src_run.font.size
        if src_run.font.color and src_run.font.color.rgb:
            body_run.font.color.rgb = src_run.font.color.rgb
    body_run.bold = False
    body_run.italic = False
    body_run.underline = False

    # Copy alignment and margins from the previous clause (after_para)
    new_para.alignment = after_para.alignment
    if after_para.paragraph_format.left_indent is not None:
        new_para.paragraph_format.left_indent = after_para.paragraph_format.left_indent
    if after_para.paragraph_format.right_indent is not None:
        new_para.paragraph_format.right_indent = after_para.paragraph_format.right_indent
    if after_para.paragraph_format.first_line_indent is not None:
        new_para.paragraph_format.first_line_indent = after_para.paragraph_format.first_line_indent

    # Renumber all top-level numbered clauses, preserving formatting
    renumber_clauses_preserve_formatting(document)

    return new_para

def insert_styled_clause_before_clause(document, before_clause_heading, new_clause_heading, new_clause_body):
    """
    Inserts a new clause (as a single paragraph: number, heading, and body together) before the specified clause heading (by text, not by section),
    and renumbers all top-level numbered clauses in the document to maintain sequential numbering.
    - before_clause_heading: the text of the clause heading before which to insert (e.g., '10. Prohibition on Use of Open AI Systems.' or 'Prohibition on Use of Open AI Systems.')
    - new_clause_heading: the heading text for the new clause (without number)
    - new_clause_body: the body text for the new clause
    Returns the new clause paragraph.
    """
    def normalize_heading(text):
        # Remove leading number and period, then strip and lowercase
        return re.sub(r'^\d+\.\s*', '', text.strip()).lower()

    def extract_clause_heading(text):
        match = re.match(r'^(\d+\.\s*[^.]+\.)', text.strip())
        if match:
            return match.group(1)
        return text.strip()

    norm_target = normalize_heading(before_clause_heading)
    before_index = None
    for i, para in enumerate(document.paragraphs):
        para_heading = extract_clause_heading(para.text)
        if normalize_heading(para_heading) == norm_target:
            before_index = i
            break
    if before_index is None:
        raise ValueError(f"Clause heading '{before_clause_heading}' not found in document.")

    # Insert a blank paragraph before the target clause, then insert the new clause after that blank paragraph
    before_para = document.paragraphs[before_index]
    blank_para = before_para.insert_paragraph_before("")
    new_para = insert_paragraph_after(blank_para, "")

    # Robustly split number, heading, and period based on the clause being inserted before
    before_text = before_para.text.strip()
    first_dot_space = before_text.find('. ')
    last_dot = before_text.rfind('.')
    if first_dot_space != -1 and last_dot != -1 and last_dot > first_dot_space:
        number_part = before_text[:first_dot_space+2]  # include '. '
        period_part = before_text[last_dot] + " "
    else:
        number_part = "XXX. "
        period_part = ". "
    heading_part = new_clause_heading

    style_info = detect_clause_heading_style(document)
    src_run = before_para.runs[0] if before_para.runs else None
    add_heading_runs(new_para, number_part, heading_part, period_part, src_run, style_info)
    # Add the body (not bold/underlined)
    body_run = new_para.add_run(new_clause_body)
    if src_run:
        if src_run.font.name:
            body_run.font.name = src_run.font.name
        if src_run.font.size:
            body_run.font.size = src_run.font.size
        if src_run.font.color and src_run.font.color.rgb:
            body_run.font.color.rgb = src_run.font.color.rgb
    body_run.bold = False
    body_run.italic = False
    body_run.underline = False

    # Copy alignment and margins from the clause being inserted before
    new_para.alignment = before_para.alignment
    if before_para.paragraph_format.left_indent is not None:
        new_para.paragraph_format.left_indent = before_para.paragraph_format.left_indent
    if before_para.paragraph_format.right_indent is not None:
        new_para.paragraph_format.right_indent = before_para.paragraph_format.right_indent
    if before_para.paragraph_format.first_line_indent is not None:
        new_para.paragraph_format.first_line_indent = before_para.paragraph_format.first_line_indent

    # Renumber all top-level numbered clauses, preserving formatting
    renumber_clauses_preserve_formatting(document)

    return new_para

def insert_sentence_in_clause(doc, starts_with, sentence, index):
    """
    Inserts a sentence into the paragraph whose text starts with 'starts_with'.
    - doc: python-docx Document
    - starts_with: string that the target paragraph should start with
    - sentence: the sentence to insert (should end with punctuation)
    - index: the index at which to insert the sentence (-1 for last)
    Returns the modified paragraph (or raises if not found).
    """
    target_para = None
    for para in doc.paragraphs:
        if para.text.strip().lower().startswith(starts_with.lower()):
            target_para = para
            break
    if target_para is None:
        raise ValueError(f"No paragraph starts with '{starts_with}'")

    # Split into sentences (naive split)
    sentence_end_re = re.compile(r'([.!?])\s+')
    text = target_para.text.strip()
    parts = sentence_end_re.split(text)
    sentences = []
    for i in range(0, len(parts)-1, 2):
        sentences.append(parts[i] + parts[i+1].strip())
    if len(parts) % 2 == 1 and parts[-1].strip():
        sentences.append(parts[-1].strip())

    # Insert the new sentence
    if index == -1 or index >= len(sentences):
        sentences.append(sentence.strip())
    else:
        sentences.insert(index, sentence.strip())

    # Rebuild the paragraph text
    new_text = ' '.join(sentences)
    # Detect if original paragraph was ALL CAPS
    was_all_caps = all(
        (getattr(run.font, 'all_caps', None) or run.text.isupper())
        for run in target_para.runs if run.text.strip()
    )
    # Remove all runs from the paragraph
    for run in target_para.runs:
        run.text = ''
    # Add a new run with the new text
    new_run = target_para.add_run(new_text)
    if was_all_caps:
        new_run.font.all_caps = True
    return target_para
