import os
from docx import Document
from docx.oxml.ns import qn
from utils import (
    insert_styled_numbered_clause_after_heading,
    insert_styled_clause_after_clause,
    insert_styled_clause_before_clause,
    insert_sentence_in_clause,
)


# --- Configurable parameters ---

OUTPUT_DIR = 'updated_contracts'

# --- Contract 1 ---
INPUT_FILE_1 = 'Contract 1.docx'
OUTPUT_FILE_1 = os.path.join(OUTPUT_DIR, INPUT_FILE_1)

# The clause to insert
NEW_CLAUSE_HEADING_1 = 'Affiliate'
NEW_CLAUSE_BODY_1 = (
    'means any entity that directly or indirectly controls, is controlled by, or is under common control with a party, '
    'where "control" means the possession, directly or indirectly, of the power to direct or cause the direction of the management '
    'and policies of such entity, whether through ownership of voting securities, by contract, or otherwise.'
)
INSERT_AFTER_HEADING_1 = 'Definitions.'

# --- Ensure output directory exists ---
os.makedirs(OUTPUT_DIR, exist_ok=True)

# --- Load document ---
doc_1 = Document(INPUT_FILE_1)

# --- Insert the clause using the new utility function ---
# To insert after a specific clause within the section, use:
# insert_styled_numbered_clause_after_heading(doc, INSERT_AFTER_HEADING, NEW_CLAUSE_HEADING, NEW_CLAUSE_BODY, after_clause="Some Clause Heading")
insert_styled_numbered_clause_after_heading(
    doc_1,
    INSERT_AFTER_HEADING_1,
    NEW_CLAUSE_HEADING_1,
    NEW_CLAUSE_BODY_1,
    # after_clause="Authorized Users"
)

# --- Save updated document ---
doc_1.save(OUTPUT_FILE_1)
print(f"Updated contract saved to {OUTPUT_FILE_1}")



# --- Contract 2 ---
INPUT_FILE_2 = 'Contract 2.docx'
OUTPUT_FILE_2 = os.path.join(OUTPUT_DIR, INPUT_FILE_2)

STARTS_WITH_2 = 'The Disclosing Party is providing Confidential Information on an “as is” basis'
# The sentence to insert
NEW_SENTENCE_2 = (
    'The Disclosing Party makes no representations or warranties regarding the accuracy or completeness of the Confidential Information.'
)

# --- Ensure output directory exists ---
os.makedirs(OUTPUT_DIR, exist_ok=True)

# --- Load document ---
doc_2 = Document(INPUT_FILE_2)


insert_sentence_in_clause(
    doc_2,
    STARTS_WITH_2,
    NEW_SENTENCE_2,
    1,
)

# --- Save updated document ---
doc_2.save(OUTPUT_FILE_2)
print(f"Updated contract saved to {OUTPUT_FILE_2}")



# --- Contract 3 ---
INPUT_FILE_3 = 'Contract 3.docx'
OUTPUT_FILE_3 = os.path.join(OUTPUT_DIR, INPUT_FILE_3)

# The clause to insert
NEW_CLAUSE_HEADING_3 = 'Residuals'
NEW_CLAUSE_BODY_3 = (
    'Nothing in this Agreement shall be construed to limit the Receiving Party\'s right to '
    'independently develop or acquire products or services without use of the Disclosing Party\'s '
    'Confidential Information, nor shall it restrict the use of any general knowledge, skills, or '
    'experience retained in unaided memory by personnel of the Receiving Party.'
)
INSERT_AFTER_HEADING_3 = None

# --- Ensure output directory exists ---
os.makedirs(OUTPUT_DIR, exist_ok=True)

# --- Load document ---
doc_3 = Document(INPUT_FILE_3)

# --- Insert the clause using the new utility function ---
# To insert after a specific clause within the section, use:
# insert_styled_numbered_clause_after_heading(doc, INSERT_AFTER_HEADING, NEW_CLAUSE_HEADING, NEW_CLAUSE_BODY, after_clause="Some Clause Heading")
insert_styled_clause_after_clause(
    doc_3,
    "Prohibition on Use of Open AI Systems.",
    NEW_CLAUSE_HEADING_3,
    NEW_CLAUSE_BODY_3,
)

# --- Save updated document ---
doc_3.save(OUTPUT_FILE_3)
print(f"Updated contract saved to {OUTPUT_FILE_3}")
